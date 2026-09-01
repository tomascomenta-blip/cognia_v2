# -*- coding: utf-8 -*-
"""
cognia/clases/vista.py
======================
EL CUADERNO VIRTUAL: un solo HTML donde el duenio ve todo lo que ha dado en
clase, por materia, y lo lee como un cuaderno de verdad.

POR QUE UN FICHERO SUELTO Y NO UN SERVIDOR. Es el patron de vista de la casa
(memory/memorias_view.py, agent/flujoteca_view.py): se escribe a ~/.cognia/ y
se abre con webbrowser, sin CDN ni red. Aqui ademas ese formato es el PRODUCTO
y no una comodidad: el cuaderno de un curso se manda por correo, se guarda en
un pendrive y se abre dentro de dos anios en otro ordenador. Por eso las
imagenes y los clips viajan EMBEBIDOS como data: URI -- un <img src="fisica/
pizarra_0003.png"> convierte el cuaderno en una carpeta que se rompe en cuanto
alguien mueve un fichero.

EL LIMITE DE ESO, DECLARADO. base64 engorda 4/3, y una jornada con 40 fotos de
pizarra a 3 MB seria un HTML de 160 MB que ningun navegador abre comodo. Hay
dos topes (TOPE_ADJUNTO por fichero y TOPE_TOTAL para la pagina entera) y lo
que no cabe NO desaparece: se enlaza con file:// y la propia pagina dice, en
esa entrada, que ese adjunto no viaja con el HTML y por que. El fallo tipico de
este repo es el vacio silencioso; una foto que falta y no se explica es
exactamente eso.

Los dos topes se cuentan en unidades DISTINTAS y a proposito. TOPE_ADJUNTO
mira el fichero original, porque es el numero del que el duenio habla ("una
foto de 3 MB"). TOPE_TOTAL mira lo que ese adjunto ANIADE AL HTML, que es 4/3
del original: es un presupuesto de peso de pagina, y contarlo en bytes de
origen dejaba pasar un fichero un 35% mas gordo que el tope y hacia que el pie
anunciara menos peso del que la pagina tiene. Ese pie existe justo para que el
duenio sepa si el correo va a rebotar, asi que ahi no puede haber un 35% de
diferencia entre lo que dice y lo que pesa.

APUNTES vs TRANSCRIPCION. El cuaderno son los APUNTES; la transcripcion es la
FUENTE. Por eso el resumen, las claves, las formulas y lo que entra en examen
se ven abiertos y la transcripcion entera va en un <details> plegado. Aun asi
la transcripcion cuenta para el buscador: el duenio busca "efecto Doppler" y
tiene que caer la clase donde el profesor lo dijo, aunque no lo apuntara nadie.

SEGURIDAD. Todo lo que se pinta viaja como DATO en un JSON embebido y se mete
en el DOM con textContent o con setAttribute sobre plantillas <template>: la
pagina no tiene ni un innerHTML. Y el JSON va con TODOS los "<" escapados como
\\u003c, no solo el "</". El motivo NO es una medida de este repo sino lo que
dice el estandar de tokenizacion de HTML (WHATWG 13.2.5.15 y siguientes):
"<!--" y "<script" hacen entrar al tokenizador en 'script data escaped', y en
ese estado el </script> de la plantilla ya no cierra el bloque -- se traga el
resto del documento y la pagina queda muda. Una nota de clase que copie codigo
HTML es el caso NORMAL, no un ataque.

Y con "<" no basta: json.dumps(ensure_ascii=False) deja crudos U+2028 y U+2029,
que son TERMINADORES DE LINEA para JavaScript. Dentro de un literal de cadena
eso es un error de sintaxis en cualquier motor anterior a ES2019, y el sintoma
es el mismo: la pagina entera muda. No es rebuscado -- U+2028 aparece solo al
pegar texto de un PDF, que es de donde salen la mitad de los apuntes. Se
escapan los tres en _escapar_para_script().

UN CUADERNO POR ASIGNATURA. El duenio lo pidio asi: "que cada materia se
guarde en un cuaderno distinto, para que no se mezclen todas las materias".
`export_materias()` escribe un HTML por materia mas un `indice.html` que
enlaza a todos. No es solo comodidad de lectura: el PRESUPUESTO de pagina
(TOPE_TOTAL) se gasta POR FICHERO, asi que separar por materia hace que las
20 fotos de pizarra de Fisica ya no dejen sin imagen a las clases de Historia
que se pintaban despues (ver `_embeber`: el presupuesto se gasta por orden, y
lo que no cabe cae a enlace file://). Cada fichero dice en su pie lo que ESE
cuaderno pesa, y el indice repite el peso de cada uno.

COMO SALE DE AQUI A PAPEL Y A UN PROCESADOR DE TEXTOS. Tres caminos, y el
primero es el bueno:

  1. PDF por el navegador (UNIVERSAL, sin instalar nada): boton "Imprimir" ->
     "Guardar como PDF". La hoja de estilo tiene un bloque @media print que
     quita cabecera, nav y pie, y evita que una sesion, una ficha, una foto o
     una formula se parta entre dos hojas. Antes de imprimir, la pagina llama
     a `__prepararImpresion()` desde el evento `beforeprint`: abre las
     transcripciones plegadas (un <details> cerrado imprime SOLO su titulo),
     pone las imagenes en loading="eager" (un seguro cuyo alcance real esta
     medido en `_preparar_papel`: hoy no cambia nada) y aniade al final la
     nota de cuantos clips de audio se quedan fuera del papel -- en un PDF no
     suena ninguno, y eso hay que decirlo, no dejarlo notar.
  2. `export_pdf()` con playwright: el MISMO camino automatizado, para generar
     el PDF sin abrir el navegador. Es un EXTRA OPCIONAL y hay que decir por
     que: playwright NO esta en el venv del producto (~/.cognia/venv), asi que
     este camino corre en el repo y falla en una instalacion limpia. Cuando
     falta, el error dice los DOS pasos (pip install playwright Y playwright
     install chromium) en vez de fingir que es uno.
  3. `export_dom()` para subir a un procesador de textos. MEDIDO: subir el
     HTML crudo a Google Docs da un documento VACIO de 272 caracteres, porque
     el contenido lo pinta el JS al abrir y el importador no ejecuta JS. Lo
     que se sube es el DOM YA RENDERIZADO (page.content() despues del JS), que
     si se convierte. Y si hay python-docx, `export_docx()` escribe un .docx
     directo desde los bloques del documento de la materia.

Sin modelo. Esta vista no llama al LLM: pinta lo que ya hay en el cuaderno.
Los apuntes (titulo, resumen, claves) los produce quien los escribio en
apuntes.json; aqui si no estan, la sesion se ve igual con su linea de tiempo.
"""

from __future__ import annotations

import base64
import html as _html
import json
import logging
import re
import time
from contextlib import contextmanager
from pathlib import Path

from cognia.clases import almacen as alm
from cognia.clases import cuaderno as cua

log = logging.getLogger(__name__)

__all__ = ["render_html", "construir", "export", "export_materias",
           "export_pdf", "export_dom", "export_docx", "nombre_de_fichero",
           "FICHERO_INDICE", "ErrorExportacion",
           "TOPE_ADJUNTO", "TOPE_TOTAL"]


class ErrorExportacion(RuntimeError):
    """Una exportacion que NO se pudo hacer, con el motivo y que hacer.

    Lleva siempre los pasos exactos en el mensaje: el llamante tipico es el
    CLI, y "no pude generar el PDF" sin decir que falta obliga al duenio a
    adivinar entre 'no esta instalado', 'esta a medias' y 'se rompio'.
    """

# Topes de embebido. DEFAULTS CONSERVADORES, NO MEDIDOS: nadie ha cronometrado
# todavia en esta maquina a partir de que peso un navegador se atraganta con un
# data: URI, ni cuanto pesa la foto de pizarra tipica del duenio. 4 MB por
# adjunto y 64 MB de pagina son el orden de magnitud que deja pasar una jornada
# entera de fotos sin llegar a los limites de los que si hay constancia publica
# (Chrome corta un data: URI de navegacion en 2 MB, pero como src de <img> no
# documenta tope). En cuanto haya una medida real se cambian aqui.
#
# UNIDADES, que no son la misma en los dos (ver cabecera):
#   TOPE_ADJUNTO -> bytes del fichero ORIGINAL en disco.
#   TOPE_TOTAL   -> bytes que la pagina ENGORDA, o sea el largo del data: URI.
TOPE_ADJUNTO = 4 * 1024 * 1024
TOPE_TOTAL = 64 * 1024 * 1024

# Extensiones que se embeben, con su MIME. Es una lista CERRADA a proposito:
# el src de un <img>/<audio> se construye con esto, y adivinar el MIME con
# mimetypes (que en Windows sale del registro y varia entre maquinas) haria
# que el mismo cuaderno se viera distinto en dos ordenadores.
_MIME_IMAGEN = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                ".gif": "image/gif", ".webp": "image/webp", ".bmp": "image/bmp"}
_MIME_AUDIO = {".wav": "audio/wav", ".mp3": "audio/mpeg", ".ogg": "audio/ogg",
               ".oga": "audio/ogg", ".m4a": "audio/mp4", ".mp4": "audio/mp4",
               ".flac": "audio/flac", ".webm": "audio/webm"}

# Etiqueta legible por tipo de entrada. El duenio no tiene por que saber que
# 'referencia' es un tipo del modelo de datos.
_ETIQUETA_TIPO = {
    cua.TIPO_NOTA: "Nota",
    cua.TIPO_IMAGEN: "Imagen",
    cua.TIPO_AUDIO: "Clip",
    cua.TIPO_REFERENCIA: "Referencia",
    cua.TIPO_MARCA: "Marca",
}

_DIAS = ("lunes", "martes", "miercoles", "jueves", "viernes", "sabado", "domingo")

# Como se leen los apuntes. Los escribe OTRO modulo (apuntes.json), asi que
# aqui se aceptan varias ortografias por campo: si manianas alguien guarda
# "puntos_clave" en vez de "claves", una tabla de un solo nombre dejaria la
# ficha vacia SIN decir nada -- otra vez el vacio silencioso. Y lo que no
# encaje en ninguna fila no se tira: cae en 'otros' y se ve igual.
_CAMPOS_APUNTES = (
    ("titulo", ("titulo", "title"), "texto"),
    ("resumen", ("resumen", "sumario", "summary"), "texto"),
    ("claves", ("claves", "puntos_clave", "puntos", "clave", "ideas"), "lista"),
    ("definiciones", ("definiciones", "conceptos", "vocabulario"), "lista"),
    ("formulas", ("formulas", "ecuaciones"), "lista"),
    ("deberes", ("deberes", "tareas", "ejercicios", "homework"), "lista"),
    ("examen", ("examen", "entra_en_examen", "evaluable", "para_el_examen"), "lista"),
    ("dudas", ("dudas", "preguntas"), "lista"),
)

# Titulo visible de cada ficha de apuntes, en el orden en que se pintan.
_ORDEN_FICHAS = (("claves", "Puntos clave"), ("definiciones", "Definiciones"),
                 ("formulas", "Formulas"), ("examen", "Entra en el examen"),
                 ("deberes", "Deberes"), ("dudas", "Dudas"))


# ── formato de tiempos ───────────────────────────────────────────────────────

def _duracion(segundos: float) -> str:
    """'1 h 05 min' / '42 min' / '38 s'. Se redondea a la unidad que el duenio
    usa para hablar de una clase; poner '3245 s' seria exacto e inutil."""
    s = int(max(0.0, segundos))
    if s < 60:
        return "%d s" % s
    if s < 3600:
        return "%d min" % (s // 60)
    return "%d h %02d min" % (s // 3600, (s % 3600) // 60)


def _fecha_de(jornada: cua.Jornada) -> tuple:
    """(texto legible, epoch de inicio). El epoch de la jornada manda; si
    vale 0 (una jornada importada o a medio cerrar) se cae al NOMBRE, que es
    la fecha en ISO. Preferir el nombre siempre seria peor: la jornada '-2'
    del mismo dia perderia su hora."""
    epoch = float(getattr(jornada, "inicio_epoch", 0.0) or 0.0)
    if epoch > 0:
        tm = time.localtime(epoch)
        return "%s %s" % (_DIAS[tm.tm_wday], time.strftime("%d/%m/%Y", tm)), epoch
    try:
        tm = time.strptime(jornada.nombre[:10], "%Y-%m-%d")
        return "%s %s" % (_DIAS[tm.tm_wday], time.strftime("%d/%m/%Y", tm)), 0.0
    except ValueError:
        # Nombre de jornada que no es una fecha: se ensenia crudo antes que
        # inventar un dia que el duenio no reconoceria.
        return jornada.nombre, 0.0


def _reloj(epoch: float, t: float) -> str:
    """Hora de pared de un instante de la jornada, o '' si no hay epoch. Los
    't' del cuaderno son segundos desde el inicio, no horas: sin el epoch no
    se puede inventar una hora."""
    if epoch <= 0:
        return ""
    return time.strftime("%H:%M", time.localtime(epoch + t))


def _desplazamiento(t: float) -> str:
    """'+12:34' desde el inicio de la sesion. Es lo unico que siempre existe,
    y es lo que cuadra con el audio guardado."""
    s = int(max(0.0, t))
    return "+%d:%02d:%02d" % (s // 3600, (s % 3600) // 60, s % 60) if s >= 3600 \
        else "+%02d:%02d" % (s // 60, s % 60)


# ── apuntes ──────────────────────────────────────────────────────────────────

def _lista(valor) -> list:
    """Un campo de apuntes normalizado a lista de lineas. Puede llegar como
    lista (lo normal), como parrafo con saltos, o como dict {concepto: def}
    -- las tres formas son razonables y las tres se han visto al generar
    apuntes con un modelo."""
    if valor is None or valor == "":
        return []
    if isinstance(valor, dict):
        return ["%s: %s" % (k, v) for k, v in valor.items()]
    if isinstance(valor, (list, tuple)):
        fuera = []
        for v in valor:
            if isinstance(v, dict):
                fuera += ["%s: %s" % (k, x) for k, x in v.items()]
            elif str(v).strip():
                fuera.append(str(v).strip())
        return fuera
    return [ln.strip() for ln in str(valor).splitlines() if ln.strip()]


def _leer_apuntes(crudo: dict) -> dict:
    """Los apuntes de una sesion, con los nombres de campo unificados y SIN
    perder nada: lo que no reconozco va a 'otros' con su clave original."""
    crudo = crudo if isinstance(crudo, dict) else {}
    fuera = {"titulo": "", "resumen": "", "otros": []}
    usadas = set()
    for destino, alias, forma in _CAMPOS_APUNTES:
        valor = None
        for a in alias:
            usadas.add(a)
            if a in crudo and crudo[a] not in (None, "", [], {}):
                valor = crudo[a]
                # Se para en el PRIMER alias con contenido; si por lo que sea
                # vinieran dos rellenos, el segundo no se pierde: no queda
                # marcado como usado y cae en 'otros' con su nombre.
                break
        if forma == "texto":
            fuera[destino] = str(valor).strip() if valor else ""
        else:
            fuera[destino] = _lista(valor)
    for k, v in crudo.items():
        if k in usadas or v in (None, "", [], {}):
            continue
        fuera["otros"].append({"k": str(k), "v": _lista(v)})
    return fuera


# ── adjuntos ─────────────────────────────────────────────────────────────────

def _peso_en_pagina(mime: str, tam: int) -> int:
    """Cuanto ANIADE al HTML embeber un fichero de `tam` bytes con ese MIME.

    Es exacto, no una estimacion: base64 son 4 caracteres por cada 3 bytes
    redondeando hacia arriba, mas el prefijo 'data:<mime>;base64,'. Se calcula
    ANTES de leer el fichero para no cargar 300 MB en RAM solo para descubrir
    que no caben.
    """
    return len("data:%s;base64," % mime) + 4 * ((int(tam) + 2) // 3)


def _embeber(jornada: str, nombre: str, tipo: str, gasto: dict) -> dict:
    """{'src','enlace','aviso'} para una imagen o un clip.

    Devuelve SIEMPRE algo pintable: o el data: URI, o un enlace file:// con el
    motivo por el que no viaja dentro. Ninguna rama se traga el fallo en
    silencio -- la pagina ensenia el aviso y ademas queda en el log.
    """
    fuera = {"src": "", "enlace": "", "aviso": "", "bytes": 0}
    ruta = alm.ruta_adjunto(jornada, nombre)   # _seguro() ya impide salir de la carpeta
    if not ruta.is_file():
        fuera["aviso"] = "el adjunto '%s' ya no esta en disco" % nombre
        log.warning("clases.vista: adjunto ausente %s", ruta)
        return fuera
    tabla = _MIME_IMAGEN if tipo == cua.TIPO_IMAGEN else _MIME_AUDIO
    mime = tabla.get(ruta.suffix.lower())
    try:
        tam = ruta.stat().st_size
    except OSError as exc:
        fuera["aviso"] = "no pude medir '%s': %s" % (nombre, exc)
        log.warning("clases.vista: stat fallo en %s: %s", ruta, exc)
        return fuera
    fuera["enlace"] = ruta.as_uri()
    if mime is None:
        fuera["aviso"] = ("'%s' no es un formato que sepa embeber (%s); queda "
                          "como enlace a este ordenador" % (nombre, ruta.suffix or "sin extension"))
        return fuera
    if tam > TOPE_ADJUNTO:
        fuera["aviso"] = ("'%s' pesa %.1f MB y el tope por adjunto son %.0f MB: "
                          "no viaja dentro del HTML, solo el enlace"
                          % (nombre, tam / 1048576.0, TOPE_ADJUNTO / 1048576.0))
        return fuera
    if gasto["usado"] + _peso_en_pagina(mime, tam) > TOPE_TOTAL:
        fuera["aviso"] = ("la pagina ya pesa %.1f MB de adjuntos (tope %.0f MB): "
                          "'%s' queda como enlace"
                          % (gasto["usado"] / 1048576.0, TOPE_TOTAL / 1048576.0, nombre))
        return fuera
    try:
        crudo = ruta.read_bytes()
    except OSError as exc:
        fuera["aviso"] = "no pude leer '%s': %s" % (nombre, exc)
        log.warning("clases.vista: lectura fallo en %s: %s", ruta, exc)
        return fuera
    fuera["src"] = "data:%s;base64,%s" % (mime, base64.b64encode(crudo).decode("ascii"))
    # El presupuesto se gasta con lo que REALMENTE se aniade al HTML, no con el
    # tamanio del fichero de origen: son cosas distintas por un factor 4/3, y
    # ese factor es el que el pie de la pagina le promete al duenio que no hay.
    fuera["bytes"] = len(fuera["src"])
    gasto["usado"] += fuera["bytes"]
    fuera["enlace"] = ""   # esta dentro: el enlace al disco solo confundiria
    return fuera


# ── construccion de los datos ────────────────────────────────────────────────

def _sesion_a_dict(s: cua.Sesion, jor: cua.Jornada, gasto: dict, avisos: list) -> dict:
    fecha, epoch = _fecha_de(jor)
    ap = _leer_apuntes(s.apuntes)
    linea, dicho, busca = [], [], [s.materia, jor.nombre, fecha]

    for e in s.entradas:
        rel = max(0.0, e.t - s.t0)
        if e.tipo == cua.TIPO_TRANSCRIPCION:
            if e.texto:
                dicho.append({"hora": _reloj(epoch, e.t), "marca": _desplazamiento(rel),
                              "texto": e.texto})
                busca.append(e.texto)
            continue
        item = {"tipo": e.tipo, "etiqueta": _ETIQUETA_TIPO.get(e.tipo, e.tipo),
                "hora": _reloj(epoch, e.t), "marca": _desplazamiento(rel),
                "texto": e.texto, "importante": bool(e.importante),
                "fuente": e.fuente, "src": "", "enlace": "", "aviso": ""}
        if e.adjunto and e.tipo in (cua.TIPO_IMAGEN, cua.TIPO_AUDIO):
            medio = _embeber(s.jornada or jor.nombre, e.adjunto, e.tipo, gasto)
            item.update({k: medio[k] for k in ("src", "enlace", "aviso")})
            item["adjunto"] = e.adjunto
            if medio["aviso"]:
                avisos.append("%s · %s" % (s.materia, medio["aviso"]))
        busca.append(e.texto)
        linea.append(item)

    busca += [ap["titulo"], ap["resumen"]]
    for clave, _ in _ORDEN_FICHAS:
        busca += ap[clave]
    for otro in ap["otros"]:
        busca += otro["v"]

    return {
        "materia": s.materia, "jornada": s.jornada or jor.nombre,
        "fecha": fecha, "hora": _reloj(epoch, s.t0), "hora_fin": _reloj(epoch, s.t1),
        "duracion": _duracion(s.duracion), "segundos": s.duracion,
        "confianza": round(float(s.confianza or 0.0), 2), "por": s.por,
        "apuntes": ap, "linea": linea, "dicho": dicho,
        "n_dicho": len(dicho),
        # El heno del buscador se calcula AQUI y no en JS: asi la busqueda
        # alcanza la transcripcion aunque este plegada, y cuesta lo mismo
        # buscar en un cuaderno de un dia que en el de un curso entero.
        "busca": " ".join(x for x in busca if x).lower(),
    }


def _sello(ahora=None) -> str:
    """La fecha de generacion que va en el pie, con el instante INYECTABLE.

    Mismo patron que `olvido._epoch(ahora)`, y por el mismo motivo: sin esto
    render_html() no es una funcion de sus datos -- el mismo cuaderno da un
    HTML distinto cada minuto, y ningun test puede fijar la pagina entera.
    Acepta None (reloj de pared), un epoch o un datetime.
    """
    if ahora is None:
        epoch = time.time()
    elif hasattr(ahora, "timestamp"):
        epoch = float(ahora.timestamp())
    else:
        epoch = float(ahora)
    return time.strftime("%d/%m/%Y %H:%M", time.localtime(epoch))


def construir(materias=None, ahora=None, agrupado=None) -> dict:
    """El dict que se embebe en la pagina: el cuaderno entero, por materia.

    `materias` es una lista de nombres para filtrar (se la pasa tal cual a
    cuaderno.cuaderno). None = todo. `ahora` fija el sello de generacion
    (None = reloj de pared).

    `agrupado` es el {materia: [Sesion]} YA LEIDO, para no volver a leerlo.
    Existe por el cuaderno por asignatura: exportar 10 materias son 10
    llamadas, y sin esto cada una vuelve a abrir las jornadas de esa materia.
    MEDIDO en un curso sintetico de 180 jornadas y 10 asignaturas (ver el
    bloque de medidas de `export_materias`): el bucle de 10 materias baja de
    1098 ms a 389 ms. Lo que NO comparte es el presupuesto de adjuntos:
    `gasto` se crea en cada llamada, asi que cada fichero tiene su TOPE_TOTAL
    entero -- que es justamente el motivo por el que partir el cuaderno mejora
    el reparto de fotos.
    """
    # monotonic y no time(): 'ms' es una DURACION, y con time() un ajuste de
    # NTP o el cambio de hora en mitad de una exportacion larga la deja en
    # negativo.
    t0 = time.monotonic()
    generado = _sello(ahora)
    avisos: list = []
    gasto = {"usado": 0}
    try:
        if agrupado is None:
            agrupado = cua.cuaderno(materias)
        elif materias:
            # El filtro se vuelve a aplicar aunque venga masticado: quien pasa
            # `agrupado` no tiene por que haberlo filtrado, y un cuaderno de
            # Fisica con una sesion de Historia dentro es exactamente lo que
            # el duenio pidio que no pasara.
            agrupado = {m: s for m, s in agrupado.items() if m in materias}
    except Exception as exc:
        # La vista tiene que ABRIR aunque el cuaderno este roto: es justo la
        # herramienta a la que el duenio va cuando algo no cuadra. Se abre
        # vacia DICIENDO por que, que no es lo mismo que abrirse vacia.
        log.warning("clases.vista: no pude leer el cuaderno: %s", exc)
        return {"materias": [], "total_sesiones": 0, "bytes_embebidos": 0,
                "avisos": ["no pude leer el cuaderno: %s: %s"
                           % (type(exc).__name__, exc)],
                "generado": generado, "ms": 0}

    jornadas: dict = {}   # cache: una jornada la comparten todas sus sesiones
    materias_out = []
    total = 0
    for nombre_materia in sorted(agrupado, key=lambda m: m.lower()):
        sesiones = []
        for s in agrupado[nombre_materia]:
            clave = s.jornada
            if clave not in jornadas:
                try:
                    jornadas[clave] = cua.cargar_jornada(clave)
                except Exception as exc:
                    log.warning("clases.vista: jornada %s ilegible: %s", clave, exc)
                    avisos.append("no pude leer el estado de la jornada %s: %s"
                                  % (clave, exc))
                    jornadas[clave] = cua.Jornada(nombre=clave)
            sesiones.append(_sesion_a_dict(s, jornadas[clave], gasto, avisos))
        segundos = sum(x["segundos"] for x in sesiones)
        total += len(sesiones)
        materias_out.append({
            "nombre": nombre_materia, "n": len(sesiones),
            "segundos": segundos, "horas": _duracion(segundos),
            "sesiones": sesiones,
        })
    # La materia con mas horas primero: es la que el duenio abre mas veces.
    materias_out.sort(key=lambda m: (-m["segundos"], m["nombre"].lower()))
    return {"materias": materias_out, "total_sesiones": total,
            "bytes_embebidos": gasto["usado"], "avisos": avisos,
            "generado": generado,
            "ms": int((time.monotonic() - t0) * 1000)}


# ── la pagina ────────────────────────────────────────────────────────────────

_HTML = r"""<!doctype html><html lang="es"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>__TITULO__</title>
<style>
/* Claro por defecto (un cuaderno se lee sobre papel) y oscuro cuando el
   sistema lo pide. Los tokens se REDEFINEN, no se invierten: los contrastes
   estan elegidos uno a uno en los dos modos. El boton de tema manda sobre el
   sistema y se recuerda en localStorage. */
:root{
  --fondo:#fbfaf7; --papel:#ffffff; --panel:#f2f1ec; --borde:#dcd9d0;
  --texto:#1f2328; --texto2:#6b6a63; --acento:#0969da; --acento2:#0550ae;
  --marca:#9a6700; --marcafondo:#fff4d6; --sombra:0 1px 3px rgba(31,35,40,.10);
}
@media (prefers-color-scheme: dark){
  :root:not([data-tema="claro"]){
    --fondo:#0d1117; --papel:#161b22; --panel:#1c2128; --borde:#30363d;
    --texto:#e6edf3; --texto2:#9198a1; --acento:#58a6ff; --acento2:#79c0ff;
    --marca:#e3b341; --marcafondo:#332a10; --sombra:0 1px 3px rgba(0,0,0,.45);
  }
}
:root[data-tema="oscuro"]{
  --fondo:#0d1117; --papel:#161b22; --panel:#1c2128; --borde:#30363d;
  --texto:#e6edf3; --texto2:#9198a1; --acento:#58a6ff; --acento2:#79c0ff;
  --marca:#e3b341; --marcafondo:#332a10; --sombra:0 1px 3px rgba(0,0,0,.45);
}
*{box-sizing:border-box}
html,body{margin:0;height:100%}
body{background:var(--fondo);color:var(--texto);display:flex;flex-direction:column;
  font:15px/1.6 -apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,Helvetica,Arial,sans-serif;}
header{display:flex;align-items:center;gap:14px;flex-wrap:wrap;padding:12px 20px;
  background:var(--panel);border-bottom:1px solid var(--borde);position:sticky;top:0;z-index:10}
h1{margin:0;font-size:16px;font-weight:600;letter-spacing:-.01em}
h1 span{color:var(--texto2);font-weight:400;font-size:13px;margin-left:8px}
#buscar{flex:1;min-width:200px;max-width:480px;padding:7px 12px;border-radius:6px;
  border:1px solid var(--borde);background:var(--papel);color:var(--texto);font-size:14px}
#buscar:focus{outline:none;border-color:var(--acento)}
button.btn{padding:7px 12px;border-radius:6px;border:1px solid var(--borde);
  background:var(--papel);color:var(--texto);font-size:13px;cursor:pointer}
button.btn:hover{border-color:var(--acento);color:var(--acento)}
#avisos{padding:8px 20px;background:var(--marcafondo);color:var(--texto);
  border-bottom:1px solid var(--borde);font-size:13px}
#avisos ul{margin:6px 0 0;padding-left:20px}
main{flex:1;display:flex;min-height:0}
nav{width:230px;flex:0 0 230px;overflow-y:auto;padding:14px 10px;
  background:var(--panel);border-right:1px solid var(--borde)}
nav h2{margin:0 0 8px 8px;font-size:11px;font-weight:600;letter-spacing:.06em;
  text-transform:uppercase;color:var(--texto2)}
.mat{display:block;width:100%;text-align:left;margin-bottom:3px;padding:7px 10px;
  border:1px solid transparent;border-radius:7px;background:none;color:var(--texto);
  font:inherit;cursor:pointer}
.mat:hover{background:var(--papel)}
.mat[aria-pressed="true"]{background:var(--papel);border-color:var(--acento);color:var(--acento)}
a.mat{text-decoration:none}
a.mat[aria-current="page"]{border-color:var(--acento);color:var(--acento)}
nav h2.otros{margin-top:14px}
.mat .mn{display:block;font-weight:600;font-size:14px}
.mat .mc{display:block;font-size:12px;color:var(--texto2)}
#hojas{flex:1;overflow-y:auto;padding:18px 24px 60px}
.sesion{background:var(--papel);border:1px solid var(--borde);border-radius:10px;
  box-shadow:var(--sombra);padding:16px 18px;margin:0 auto 16px;max-width:900px}
.sesion h3{margin:0 0 2px;font-size:17px;line-height:1.35}
.meta{display:flex;flex-wrap:wrap;gap:12px;color:var(--texto2);font-size:12.5px;
  font-variant-numeric:tabular-nums;margin-bottom:10px}
.resumen{margin:0 0 12px;font-size:15px}
.fichas{display:grid;gap:10px;grid-template-columns:repeat(auto-fit,minmax(240px,1fr));margin-bottom:14px}
.ficha{border:1px solid var(--borde);border-radius:8px;padding:8px 12px;background:var(--fondo)}
.ficha h5{margin:0 0 4px;font-size:11px;font-weight:600;letter-spacing:.05em;
  text-transform:uppercase;color:var(--texto2)}
.ficha ul{margin:0;padding-left:18px;font-size:14px}
.ficha.examen{border-color:var(--marca);background:var(--marcafondo)}
h4.th{margin:14px 0 6px;font-size:11px;font-weight:600;letter-spacing:.05em;
  text-transform:uppercase;color:var(--texto2)}
ol.linea{list-style:none;margin:0;padding:0}
.ent{display:flex;gap:12px;padding:7px 8px;border-left:2px solid var(--borde);margin-bottom:4px}
.ent .marca{flex:0 0 96px;color:var(--texto2);font-size:12px;
  font-variant-numeric:tabular-nums;white-space:nowrap}
.ent .cuerpo{flex:1;min-width:0}
.ent .tipo{display:inline-block;font-size:11px;padding:0 7px;border-radius:20px;
  border:1px solid var(--borde);color:var(--texto2);margin-right:6px}
.ent .txt{white-space:pre-wrap;word-break:break-word}
.ent.imp{border-left-color:var(--marca);background:var(--marcafondo)}
.ent.imp .tipo{border-color:var(--marca);color:var(--marca)}
.aviso{color:var(--marca);font-size:12.5px;margin-top:4px}
.aviso a{color:inherit}
img.adj{display:block;max-width:100%;height:auto;margin-top:6px;border-radius:8px;
  border:1px solid var(--borde)}
audio.adj{display:block;width:100%;max-width:420px;margin-top:6px}
details.trans{margin-top:12px;border-top:1px dashed var(--borde);padding-top:8px}
details.trans summary{cursor:pointer;color:var(--texto2);font-size:13px}
.tt{margin-top:8px;font-size:13.5px;color:var(--texto2);max-height:60vh;overflow-y:auto}
.tt p{margin:0 0 4px;display:flex;gap:10px}
.tt .m{flex:0 0 88px;font-variant-numeric:tabular-nums}
.vacio{max-width:900px;margin:40px auto;text-align:center;color:var(--texto2)}
footer{padding:7px 20px;background:var(--panel);border-top:1px solid var(--borde);
  color:var(--texto2);font-size:12px;display:flex;gap:16px;flex-wrap:wrap}
@media(max-width:820px){
  main{flex-direction:column}
  nav{width:auto;flex:0 0 auto;border-right:none;border-bottom:1px solid var(--borde);
    display:flex;flex-wrap:wrap;gap:4px}
  nav h2{display:none} .mat{width:auto}
}
/* La nota del papel: solo existe cuando se imprime. La escribe
   __prepararImpresion() y dice lo que el papel NO puede llevar (los clips de
   audio no suenan en un PDF). En pantalla estorbaria; en papel es la
   diferencia entre "faltan cosas" y "faltan estas cosas y por que". */
.notapapel{display:none}
@media print{
  header,nav,footer,#avisos{display:none!important}
  body,#hojas{display:block;overflow:visible}
  #hojas{padding:0}
  .sesion{break-inside:avoid;page-break-inside:avoid;box-shadow:none;
    border-color:#bbb;max-width:none;margin-bottom:10px}
  /* Lo de dentro tambien se protege del corte. Una sesion larga NO cabe en
     una hoja y el navegador la parte igual (break-inside:avoid solo se
     respeta si el bloque cabe): sin estas reglas el corte cae donde quiera y
     una formula o una grafica -- que aqui son IMAGENES, no texto que se
     recompone -- se imprime a medias en dos paginas. */
  .ent,.ficha,.tt p{break-inside:avoid;page-break-inside:avoid}
  img.adj{break-inside:avoid;page-break-inside:avoid;
    /* Una foto de pizarra a pantalla completa no cabe en una hoja A4 con sus
       margenes: sin el tope, la imagen sola provoca el salto que se queria
       evitar. 21 cm es la altura util de un A4 vertical menos margenes. */
    max-height:21cm;width:auto}
  audio.adj{display:none}         /* en papel no suena: lo dice .notapapel */
  .tt{max-height:none;overflow:visible}
  .notapapel{display:block;margin:14px auto 0;max-width:none;font-size:11px;
    color:#555;border-top:1px solid #bbb;padding-top:6px}
  a[href^="file:"]::after{content:" (" attr(href) ")";font-size:10px;color:#666}
}
</style></head><body>
<header>
  <h1>Cuaderno de clase <span id="sub"></span></h1>
  <input id="buscar" type="search" placeholder="Buscar en todo el cuaderno (tecla /)" autocomplete="off">
  <button class="btn" id="btnabrir" type="button">Abrir transcripciones</button>
  <button class="btn" id="btnimp" type="button">Imprimir</button>
  <button class="btn" id="btntema" type="button">Tema</button>
</header>
<div id="avisos" hidden></div>
<main>
  <nav id="materias"><h2>Materias</h2></nav>
  <div id="hojas"></div>
</main>
<footer><span id="pie"></span><span>Se graba y se corrige desde el CLI de Cognia</span></footer>
<noscript><p style="padding:20px">Este cuaderno pinta su contenido con JavaScript
(los datos van dentro del propio fichero). Abrilo en un navegador con JS activado.</p></noscript>

<!-- Plantillas. Todo lo que se pinta se clona de aqui y se rellena con
     textContent / setAttribute: en esta pagina no hay ni un innerHTML, asi que
     una nota de clase con HTML dentro se lee como texto y nunca como etiqueta. -->
<template id="t-materia"><button class="mat" type="button">
  <span class="mn"></span><span class="mc"></span></button></template>
<template id="t-sesion"><article class="sesion">
  <h3 class="tit"></h3>
  <div class="meta"><span class="fecha"></span><span class="hora"></span>
    <span class="dur"></span><span class="jor"></span><span class="conf"></span></div>
  <p class="resumen"></p>
  <div class="fichas"></div>
  <h4 class="th">Linea de tiempo</h4>
  <ol class="linea"></ol>
  <details class="trans"><summary></summary><div class="tt"></div></details>
</article></template>
<template id="t-ficha"><section class="ficha"><h5></h5><ul></ul></section></template>
<template id="t-entrada"><li class="ent"><span class="marca"></span>
  <div class="cuerpo"><div class="cab"><span class="tipo"></span></div>
    <div class="txt"></div><div class="medio"></div><div class="aviso"></div></div></li></template>
<template id="t-imagen"><img class="adj" alt="Imagen del cuaderno" loading="lazy"></template>
<template id="t-audio"><audio class="adj" controls preload="none"></audio></template>
<template id="t-dicho"><p><span class="m"></span><span class="x"></span></p></template>

<script>
const D = __DATOS__;
const $ = s => document.querySelector(s);
const tpl = id => document.getElementById(id).content.firstElementChild.cloneNode(true);

/* Tema: el sistema manda al arrancar (@media prefers-color-scheme cubre el
   caso de que el JS ni corra), el usuario manda despues. */
(function(){
  let t = null; try{ t = localStorage.getItem("cognia_cuaderno_tema"); }catch(e){}
  if(t) document.documentElement.setAttribute("data-tema", t);
})();
$("#btntema").onclick = () => {
  const actual = document.documentElement.getAttribute("data-tema");
  const sistema = window.matchMedia && window.matchMedia("(prefers-color-scheme: dark)").matches;
  const nuevo = (actual || (sistema ? "oscuro" : "claro")) === "oscuro" ? "claro" : "oscuro";
  document.documentElement.setAttribute("data-tema", nuevo);
  try{ localStorage.setItem("cognia_cuaderno_tema", nuevo); }catch(e){}
};

/* Un src solo puede ser lo que ESTE fichero genero: un data: URI o un file://
   a un adjunto. Filtrarlo aqui cuesta una linea y cierra la unica via por la
   que un dato del cuaderno llega a un atributo en vez de a un textContent. */
function urlSegura(u){ return /^(data:|file:)/i.test(u || "") ? u : ""; }

/* Un enlace a OTRO cuaderno solo puede ser un nombre de fichero .html del
   MISMO directorio: ni ruta, ni "..", ni esquema. Es el mismo criterio que
   urlSegura y por el mismo motivo -- estos nombres salen de la materia, o sea
   de algo que el duenio escribio. Se filtra por lo PROHIBIDO y no por una
   lista blanca de caracteres ASCII: "Fisica" se escribe con tilde y una lista
   blanca dejaria sin enlace justo a las asignaturas con acento. */
function ficheroSeguro(f){
  f = f || "";
  return (/\.html$/i.test(f) && !/[\\\/:]/.test(f) && f.indexOf("..") < 0) ? f : "";
}

let materiaActiva = null;

function pintarMaterias(){
  const nav = $("#materias");
  nav.querySelectorAll(".mat").forEach(b => b.remove());
  const todas = tpl("t-materia");
  todas.querySelector(".mn").textContent = "Todas";
  todas.querySelector(".mc").textContent = D.total_sesiones + " sesiones";
  todas.setAttribute("aria-pressed", materiaActiva === null);
  todas.onclick = () => { materiaActiva = null; pintarMaterias(); pintar(); };
  nav.appendChild(todas);
  D.materias.forEach(m => {
    const b = tpl("t-materia");
    b.querySelector(".mn").textContent = m.nombre;
    b.querySelector(".mc").textContent = m.n + (m.n === 1 ? " sesion · " : " sesiones · ") + m.horas;
    b.setAttribute("aria-pressed", materiaActiva === m.nombre);
    b.onclick = () => { materiaActiva = m.nombre; pintarMaterias(); pintar(); };
    nav.appendChild(b);
  });
}

/* Los OTROS cuadernos, cuando el cuaderno esta partido por asignatura. Van en
   el mismo nav que las materias: el duenio que abre "Fisica" tiene que poder
   saltar a "Historia" sin volver a la carpeta. */
function pintarEnlaces(){
  const enlaces = D.enlaces || [];
  if(!enlaces.length) return;
  const nav = $("#materias");
  const h = document.createElement("h2");
  h.className = "otros";
  h.textContent = "Otros cuadernos";
  nav.appendChild(h);
  enlaces.forEach(e => {
    const f = ficheroSeguro(e.fichero);
    if(!f) return;
    const a = document.createElement("a");
    a.className = "mat";
    a.setAttribute("href", f);
    if(e.actual) a.setAttribute("aria-current", "page");
    const n = document.createElement("span"); n.className = "mn";
    n.textContent = e.nombre; a.appendChild(n);
    const c = document.createElement("span"); c.className = "mc";
    c.textContent = e.detalle || ""; a.appendChild(c);
    nav.appendChild(a);
  });
}

/* La portada del cuaderno partido: una tarjeta por asignatura con lo que pesa
   su fichero. El peso va aqui porque es la pregunta de antes de mandar nada
   por correo, y en el indice se ven los de todas a la vez. */
function pintarIndice(){
  const cont = $("#hojas");
  cont.textContent = "";
  (D.cuadernos || []).forEach(c => {
    const art = document.createElement("article");
    art.className = "sesion";
    const h = document.createElement("h3");
    const f = ficheroSeguro(c.fichero);
    if(f){
      const a = document.createElement("a");
      a.setAttribute("href", f); a.textContent = c.nombre;
      h.appendChild(a);
    } else h.textContent = c.nombre;
    art.appendChild(h);
    const meta = document.createElement("div");
    meta.className = "meta";
    [c.detalle || "", c.peso || "", c.fichero || ""].forEach(x => {
      if(!x) return;
      const s = document.createElement("span"); s.textContent = x; meta.appendChild(s);
    });
    art.appendChild(meta);
    (c.avisos || []).forEach(a => {
      const p = document.createElement("p");
      p.className = "aviso"; p.textContent = a; art.appendChild(p);
    });
    cont.appendChild(art);
  });
  if(!(D.cuadernos || []).length){
    const v = document.createElement("div");
    v.className = "vacio";
    v.textContent = "El cuaderno esta vacio todavia: graba una clase y vuelve.";
    cont.appendChild(v);
  }
  $("#sub").textContent = (D.cuadernos || []).length + " asignaturas · " +
    D.total_sesiones + " sesiones";
  $("#pie").textContent = "indice de " + (D.cuadernos || []).length +
    " cuadernos · " + (D.peso_total || "") + " en total · generado el " + D.generado;
}

function pintarFichas(cont, ap){
  const filas = __FICHAS__;
  filas.forEach(([clave, titulo]) => {
    const v = ap[clave] || [];
    if(!v.length) return;
    const f = tpl("t-ficha");
    if(clave === "examen") f.classList.add("examen");
    f.querySelector("h5").textContent = titulo;
    const ul = f.querySelector("ul");
    v.forEach(x => { const li = document.createElement("li"); li.textContent = x; ul.appendChild(li); });
    cont.appendChild(f);
  });
  /* Un campo de apuntes que no reconozco se pinta igual con su nombre crudo:
     que apuntes.json cambie una clave no puede borrar contenido de la vista. */
  (ap.otros || []).forEach(o => {
    const f = tpl("t-ficha");
    f.querySelector("h5").textContent = o.k;
    const ul = f.querySelector("ul");
    o.v.forEach(x => { const li = document.createElement("li"); li.textContent = x; ul.appendChild(li); });
    cont.appendChild(f);
  });
}

function pintarEntrada(e){
  const li = tpl("t-entrada");
  if(e.importante) li.classList.add("imp");
  li.querySelector(".marca").textContent = (e.hora ? e.hora + "  " : "") + e.marca;
  li.querySelector(".tipo").textContent = e.etiqueta + (e.importante ? " · importante" : "");
  li.querySelector(".txt").textContent = e.texto || "";
  const medio = li.querySelector(".medio");
  const src = urlSegura(e.src);
  if(e.tipo === "imagen" && src){
    const img = tpl("t-imagen"); img.setAttribute("src", src);
    if(e.texto) img.setAttribute("alt", e.texto);
    medio.appendChild(img);
  } else if(e.tipo === "audio" && src){
    const au = tpl("t-audio"); au.setAttribute("src", src); medio.appendChild(au);
  }
  const av = li.querySelector(".aviso");
  if(e.aviso){
    av.appendChild(document.createTextNode(e.aviso));
    const en = urlSegura(e.enlace);
    if(en){
      av.appendChild(document.createTextNode(" · "));
      const a = document.createElement("a");
      a.setAttribute("href", en); a.textContent = "abrir desde el disco";
      av.appendChild(a);
    }
  }
  return li;
}

function pintarSesion(s){
  const art = tpl("t-sesion");
  art.querySelector(".tit").textContent = s.apuntes.titulo || (s.materia + " · " + s.fecha);
  art.querySelector(".fecha").textContent = s.fecha;
  art.querySelector(".hora").textContent = s.hora ? (s.hora + (s.hora_fin ? "-" + s.hora_fin : "")) : s.materia;
  art.querySelector(".dur").textContent = s.duracion;
  art.querySelector(".jor").textContent = "jornada " + s.jornada;
  /* De donde salio el corte de materia: si lo puso el detector con poca
     confianza, el duenio tiene que poder desconfiar de la etiqueta. */
  art.querySelector(".conf").textContent = s.por
    ? ("materia por " + s.por + (s.confianza ? " (" + s.confianza + ")" : "")) : "";
  const res = art.querySelector(".resumen");
  if(s.apuntes.resumen) res.textContent = s.apuntes.resumen; else res.remove();
  pintarFichas(art.querySelector(".fichas"), s.apuntes);
  const ol = art.querySelector(".linea");
  if(s.linea.length) s.linea.forEach(e => ol.appendChild(pintarEntrada(e)));
  else { art.querySelector("h4.th").remove(); ol.remove(); }
  const det = art.querySelector("details.trans");
  if(s.dicho.length){
    det.querySelector("summary").textContent =
      "Transcripcion de la clase (" + s.n_dicho + " fragmentos)";
    const tt = det.querySelector(".tt");
    s.dicho.forEach(d => {
      const p = tpl("t-dicho");
      p.querySelector(".m").textContent = d.hora || d.marca;
      p.querySelector(".x").textContent = d.texto;
      tt.appendChild(p);
    });
  } else det.remove();
  return art;
}

function pintar(){
  const q = $("#buscar").value.trim().toLowerCase();
  const cont = $("#hojas");
  cont.textContent = "";
  let n = 0;
  D.materias.forEach(m => {
    if(materiaActiva && m.nombre !== materiaActiva) return;
    m.sesiones.forEach(s => {
      if(q && s.busca.indexOf(q) < 0) return;
      cont.appendChild(pintarSesion(s)); n++;
    });
  });
  if(!n){
    const v = document.createElement("div");
    v.className = "vacio";
    v.textContent = D.total_sesiones
      ? "Nada que coincida con lo que buscaste."
      : "El cuaderno esta vacio todavia: graba una clase y vuelve.";
    cont.appendChild(v);
  }
  $("#sub").textContent = n + (n === 1 ? " sesion" : " sesiones") +
    (materiaActiva ? " de " + materiaActiva : "") + (q ? " que coinciden" : "");
}

/* Avisos arriba del todo: un adjunto que no viaja dentro del HTML tiene que
   verse ANTES de que el duenio mande el fichero por correo. */
if((D.avisos || []).length){
  const av = $("#avisos"); av.hidden = false;
  const p = document.createElement("div");
  p.textContent = "Avisos (" + D.avisos.length + "):";
  av.appendChild(p);
  const ul = document.createElement("ul");
  D.avisos.forEach(a => { const li = document.createElement("li"); li.textContent = a; ul.appendChild(li); });
  av.appendChild(ul);
}

$("#buscar").oninput = pintar;
$("#btnabrir").onclick = () => {
  const abrir = document.querySelector("details.trans:not([open])") !== null;
  document.querySelectorAll("details.trans").forEach(d => { d.open = abrir; });
  $("#btnabrir").textContent = abrir ? "Cerrar transcripciones" : "Abrir transcripciones";
};
$("#btnimp").onclick = () => window.print();
/* DEJAR LA PAGINA LISTA PARA EL PAPEL. Tres cosas que el navegador no hace
   solo, y las tres se notan como contenido que falta:
     1. Un <details> cerrado imprime SOLO su titulo -- el navegador no
        renderiza lo plegado. Se abren y se deja como estaba despues.
     2. Una <img loading="lazy"> que nunca entro en pantalla no se ha
        descargado, y al imprimir sale EN BLANCO. Pasar a "eager" la carga.
     3. En papel no suena nada: los clips de audio se quedan fuera y hay que
        DECIRLO en el pie, con su numero.
   Va en una funcion global y no dentro del listener a proposito: `page.pdf()`
   de un navegador headless NO dispara `beforeprint`, asi que export_pdf tiene
   que poder llamar exactamente a lo mismo que llama el boton Imprimir. Un
   segundo camino, aunque fuera de dos lineas, se desincronizaria del primero
   y el PDF automatico saldria distinto del que hace el duenio a mano. */
let plegadosAntes = [];
window.__prepararImpresion = function(){
  plegadosAntes = Array.from(document.querySelectorAll("details.trans:not([open])"));
  plegadosAntes.forEach(d => { d.open = true; });
  const imgs = Array.from(document.querySelectorAll("img.adj"));
  imgs.forEach(i => { i.setAttribute("loading", "eager"); });
  const clips = document.querySelectorAll("audio.adj").length;
  const enlazados = document.querySelectorAll(".aviso a").length;
  let nota = document.getElementById("papel");
  if(!nota){
    nota = document.createElement("p");
    nota.id = "papel"; nota.className = "notapapel";
    $("#hojas").appendChild(nota);
  }
  nota.textContent = "En papel: " + clips +
    (clips === 1 ? " clip de audio se queda fuera" : " clips de audio se quedan fuera") +
    " (no suenan en un PDF; se escuchan en el cuaderno HTML)" +
    (enlazados ? " · " + enlazados + " adjunto(s) no viajan dentro del fichero y solo van como enlace a este ordenador" : "") +
    " · " + imgs.length + (imgs.length === 1 ? " imagen impresa" : " imagenes impresas") + ".";
  return {clips: clips, imagenes: imgs.length, enlazados: enlazados,
          plegados: plegadosAntes.length};
};
window.addEventListener("beforeprint", () => { window.__prepararImpresion(); });
window.addEventListener("afterprint", () => { plegadosAntes.forEach(d => { d.open = false; }); });
document.addEventListener("keydown", e => {
  if(e.key === "/" && document.activeElement !== $("#buscar")){ e.preventDefault(); $("#buscar").focus(); }
  if(e.key === "Escape"){ $("#buscar").value = ""; $("#buscar").blur(); pintar(); }
});
/* El peso embebido se dice en la unidad que toca: "0.0 MB" en un cuaderno de
   7 KB parece que las fotos no viajaron, que es exactamente la duda que este
   pie existe para quitar antes de mandar el fichero por correo.
   D.bytes_embebidos son los bytes que las fotos y el audio OCUPAN EN ESTE
   HTML (el data: URI ya en base64), no los del fichero de origen: la pregunta
   que contesta el pie es "cuanto pesa esto que voy a adjuntar". */
const emb = D.bytes_embebidos < 1048576
  ? (D.bytes_embebidos / 1024).toFixed(0) + " KB"
  : (D.bytes_embebidos / 1048576).toFixed(1) + " MB";
if(D.indice){
  /* La portada del cuaderno partido por asignatura: no tiene sesiones
     propias, asi que ni buscador ni lista de materias tienen nada que hacer.
     Pintar la pagina normal aqui diria "El cuaderno esta vacio todavia"
     encima de un indice con doce asignaturas dentro. */
  $("#buscar").disabled = true;
  $("#buscar").placeholder = "Abre un cuaderno para buscar dentro";
  $("#btnabrir").remove();
  pintarIndice();
} else {
  $("#pie").textContent = D.materias.length + " materias · " + D.total_sesiones +
    " sesiones · " + emb + " de fotos y audio dentro del fichero · generado el " + D.generado;
  pintarMaterias(); pintarEnlaces(); pintar();
}
</script></body></html>"""


# Lo que un literal JSON NO puede llevar crudo dentro de un <script>, con su
# escape \uXXXX. Los tres estan dentro de una cadena JSON siempre (la sintaxis
# JSON no usa ninguno de los tres fuera de cadenas), y ahi "<" es LA MISMA
# cadena que "<" para el parser: el dato no cambia, solo deja de poder tocar al
# tokenizador de HTML ni al lexer de JavaScript.
_ESCAPES_SCRIPT = (("<", "\\u003c"),          # "</script", "<!--" -> script data escaped
                   ("\u2028", "\\u2028"),   # LINE SEPARATOR
                   ("\u2029", "\\u2029"))  # PARAGRAPH SEPARATOR
#            ^ escritos como \uXXXX en el FUENTE a proposito: un U+2028
#            literal aqui seria invisible en el editor y lo pierde cualquier
#            round-trip por una codificacion que no sea utf-8.


def _escapar_para_script(texto: str) -> str:
    """Un JSON ya serializado, listo para meter dentro de un <script>."""
    for malo, bueno in _ESCAPES_SCRIPT:
        texto = texto.replace(malo, bueno)
    return texto


def render_html(datos, titulo: str = "Cuaderno de clase", ahora=None) -> str:
    """El HTML entero, autocontenido. `datos` es lo que devuelve construir().

    `ahora` solo se usa si `datos` no trae ya su sello de generacion.
    """
    if not isinstance(datos, dict):
        raise TypeError("render_html espera el dict de construir(), no %r" % type(datos))
    datos = dict(datos)
    datos.setdefault("materias", [])
    datos.setdefault("total_sesiones", 0)
    datos.setdefault("avisos", [])
    datos.setdefault("bytes_embebidos", 0)
    datos.setdefault("generado", _sello(ahora))

    crudo = _escapar_para_script(json.dumps(datos, ensure_ascii=False))
    fichas = _escapar_para_script(json.dumps(list(_ORDEN_FICHAS), ensure_ascii=False))
    # El titulo cae dentro de <title>: ahi manda el escape de HTML, no el de JS.
    tit = _html.escape(str(titulo or "Cuaderno de clase"))
    # UNA sola pasada. Encadenar .replace() deja que lo ya sustituido se
    # reinterprete despues: un titulo que contenga "__DATOS__" se comeria el
    # JSON entero (el mismo bug que ya se pago en flujoteca_view).
    trozos = {"__TITULO__": tit, "__DATOS__": crudo, "__FICHAS__": fichas}
    return re.sub("__TITULO__|__DATOS__|__FICHAS__",
                  lambda m: trozos[m.group(0)], _HTML)


def export(path=None, open_browser: bool = True, materias=None) -> Path:
    """Escribe el cuaderno y (por defecto) lo abre. Devuelve la ruta.

    Por defecto va a ~/.cognia/cuaderno.html, como el resto de vistas de la
    casa: un sitio fijo que el duenio ya conoce y del que puede arrastrar el
    fichero a un correo.
    """
    datos = construir(materias)
    destino = Path(path).expanduser() if path else (Path.home() / ".cognia" / "cuaderno.html")
    destino.parent.mkdir(parents=True, exist_ok=True)
    sufijo = (" · " + ", ".join(materias)) if materias else ""
    destino.write_text(render_html(datos, "Cuaderno de clase" + sufijo),
                       encoding="utf-8")
    if open_browser:
        import webbrowser
        try:
            webbrowser.open(destino.as_uri())
        except Exception as exc:
            # El fichero YA esta escrito: que no haya navegador por defecto no
            # puede parecer que la exportacion fallo.
            log.warning("clases.vista: no pude abrir el navegador: %s", exc)
    return destino


# ── un cuaderno por asignatura ───────────────────────────────────────────────

FICHERO_INDICE = "indice.html"
_PREFIJO_FICHERO = "cuaderno-"


def nombre_de_fichero(materia: str, usados=None) -> str:
    """El nombre de fichero de una materia: 'cuaderno-fisica.html'.

    Pasa por `almacen._seguro` -- la MISMA sanitizacion que usan las jornadas
    y los adjuntos -- porque el nombre de la materia lo escribe el duenio y
    puede traer una barra dentro ("Fisica/Quimica" es un nombre normal de
    asignatura); dos reglas distintas de nombre seguro en el mismo cuaderno
    acabarian con dos ficheros para la misma materia.

    `usados` es el conjunto de nombres ya repartidos: dos materias distintas
    pueden sanear al MISMO fichero ("Fisica II" y "Fisica  II"), y sin
    desempate la segunda pisaria a la primera -- que es exactamente la mezcla
    que este cuaderno partido existe para evitar. El desempate es un sufijo
    numerico, no un hash: el duenio tiene que reconocer el fichero.
    """
    base = alm._seguro(materia).strip().lower().replace(" ", "-")
    base = base.strip("-") or "sin-materia"
    nombre = "%s%s.html" % (_PREFIJO_FICHERO, base)
    if usados is None:
        return nombre
    n = 2
    while nombre in usados:
        nombre = "%s%s-%d.html" % (_PREFIJO_FICHERO, base, n)
        n += 1
    usados.add(nombre)
    return nombre


def _peso(n: int) -> str:
    """Bytes en la unidad en la que el duenio decide si eso cabe en un correo."""
    return "%.0f KB" % (n / 1024.0) if n < 1048576 else "%.1f MB" % (n / 1048576.0)


def export_materias(directorio=None, materias=None, open_browser: bool = False,
                    ahora=None) -> dict:
    """UN HTML POR ASIGNATURA mas un indice que enlaza a todos.

    Lo que el duenio pidio: "que cada materia se guarde en un cuaderno
    distinto, para que no se mezclen todas las materias". Devuelve
    {'directorio', 'indice', 'ficheros': {materia: Path}, 'avisos', 'ms'}.

    EL CURSO SE LEE UNA SOLA VEZ. `cuaderno.cuaderno()` se llama con el filtro
    entero (todas las materias que se van a exportar) y el resultado se
    reparte aqui; la alternativa evidente -- una llamada por materia -- vuelve
    a abrir las jornadas de esa materia en cada vuelta.

    MEDIDO en un curso sintetico de 180 jornadas, 10 asignaturas con horario
    semanal (4 clases al dia), 48 lineas de transcripcion por jornada, 1,2 MB
    de JSONL:

        exportar las 10, una llamada por materia y SIN indice   2258 ms
        exportar las 10, una llamada por materia CON indice     1098 ms  2,1x
        exportar las 10, indice + UNA lectura del curso          389 ms  5,8x
        export_materias() entera (10 HTML + indice, escritos)    428 ms

    Y para una sola asignatura, `cuaderno(['Fisica'])`: 218 ms releyendo el
    curso entero contra 101 ms con el indice (2,2x), porque solo se abren las
    72 jornadas de las 180 en las que hay Fisica. El indice se mantiene por
    HUELLA de fichero: aniadir una jornada cuesta 24 ms UNA vez (114 ms la
    primera lectura, 90 ms las siguientes) y la sesion nueva aparece en el
    tramo sin que nadie avise al indice.

    EL PRESUPUESTO SI ES POR FICHERO, a proposito. Cada `construir()` arranca
    con TOPE_TOTAL entero, asi que las fotos de Fisica ya no dejan sin imagen
    a Historia (antes el presupuesto se gastaba por orden alfabetico sobre una
    sola pagina). El pie de cada cuaderno dice lo que ESE fichero pesa y el
    indice repite el peso de todos.
    """
    t0 = time.monotonic()
    destino = Path(directorio).expanduser() if directorio else (
        Path.home() / ".cognia" / "cuadernos")
    destino.mkdir(parents=True, exist_ok=True)
    generado = _sello(ahora)
    avisos: list = []

    if materias:
        nombres = [str(m) for m in materias]
    else:
        try:
            nombres = cua.materias_vistas()
        except Exception as exc:
            log.warning("clases.vista: no pude listar materias: %s", exc)
            avisos.append("no pude listar las materias (%s: %s): se exporta lo "
                          "que se pueda leer del cuaderno entero"
                          % (type(exc).__name__, exc))
            nombres = []
    try:
        agrupado = cua.cuaderno(nombres or None)
    except Exception as exc:
        log.warning("clases.vista: cuaderno ilegible: %s", exc)
        avisos.append("no pude leer el cuaderno: %s: %s" % (type(exc).__name__, exc))
        agrupado = {}
    # Lo que el indice NO conocia (una materia recien detectada, o el indice
    # apagado) sale igual: manda lo que hay en el cuaderno leido.
    for m in agrupado:
        if m not in nombres:
            nombres.append(m)

    usados: set = set()
    plan = [(m, nombre_de_fichero(m, usados)) for m in nombres]
    ficheros: dict = {}
    tarjetas: list = []
    total_sesiones = 0
    peso_total = 0
    for materia, fichero in plan:
        enlaces = [{"nombre": "Indice", "fichero": FICHERO_INDICE,
                    "detalle": "todas las asignaturas"}]
        enlaces += [{"nombre": otra, "fichero": f, "actual": otra == materia}
                    for otra, f in plan]
        # Se le pasa el cuaderno ENTERO y el filtro: quien separa las materias
        # es `construir`, en un solo sitio. Repartir aqui el dict a mano
        # dejaria ese filtro sin nadie que lo ejercite, y el dia que se cayera
        # el cuaderno de Fisica saldria con las clases de Historia dentro sin
        # que ningun test chillara.
        datos = construir([materia], ahora=ahora, agrupado=agrupado)
        datos["enlaces"] = enlaces
        doc = render_html(datos, "Cuaderno de clase · %s" % materia)
        ruta = destino / fichero
        try:
            ruta.write_text(doc, encoding="utf-8")
        except OSError as exc:
            # Un fichero que no se puede escribir NO se lista en el indice (no
            # se aniade su tarjeta) y el motivo sale en los avisos. Lo que si
            # puede quedar es un enlace muerto en el nav de los cuadernos ya
            # escritos, porque esos ya estan en disco: por eso el aviso va al
            # indice, que es la portada y se escribe al final.
            log.warning("clases.vista: no pude escribir %s: %s", ruta, exc)
            avisos.append("no pude escribir %s (%s: %s)"
                          % (ruta.name, type(exc).__name__, exc))
            continue
        ficheros[materia] = ruta
        n = datos["total_sesiones"]
        total_sesiones += n
        # El peso se lee del DISCO, no de len(doc.encode()): en Windows
        # write_text traduce cada "\n" a "\r\n" y el fichero acaba pesando ~119
        # bytes mas que la cadena -- suficiente para que el indice anunciara
        # "29 KB" de un fichero de 30 KB. El pie existe para que ese numero se
        # pueda creer.
        try:
            bytes_html = ruta.stat().st_size
        except OSError as exc:
            log.warning("clases.vista: no pude medir %s: %s", ruta, exc)
            avisos.append("no pude medir %s (%s): el peso que anuncia el indice "
                          "es el de la pagina en memoria" % (ruta.name, exc))
            bytes_html = len(doc.encode("utf-8"))
        peso_total += bytes_html
        avisos += ["%s · %s" % (materia, a) for a in datos["avisos"]]
        tarjetas.append({
            "nombre": materia, "fichero": fichero, "n": n,
            "detalle": "%d %s · %s" % (n, "sesion" if n == 1 else "sesiones",
                                       (datos["materias"][0]["horas"]
                                        if datos["materias"] else "0 s")),
            # El peso REAL del fichero, no el de los adjuntos: es el numero
            # que decide si el correo rebota, y el HTML pesa mas que lo que
            # lleva embebido.
            "peso": "%s en disco (%s de fotos y audio dentro)"
                    % (_peso(bytes_html), _peso(datos["bytes_embebidos"])),
            "avisos": datos["avisos"],
        })

    indice_datos = {
        "indice": True, "cuadernos": tarjetas, "materias": [],
        "total_sesiones": total_sesiones, "bytes_embebidos": 0,
        "peso_total": _peso(peso_total), "avisos": avisos, "generado": generado,
    }
    ruta_indice = destino / FICHERO_INDICE
    ruta_indice.write_text(render_html(indice_datos, "Cuadernos por asignatura"),
                           encoding="utf-8")
    if open_browser:
        import webbrowser
        try:
            webbrowser.open(ruta_indice.as_uri())
        except Exception as exc:
            log.warning("clases.vista: no pude abrir el navegador: %s", exc)
    return {"directorio": destino, "indice": ruta_indice, "ficheros": ficheros,
            "avisos": avisos, "ms": int((time.monotonic() - t0) * 1000)}


# ── PDF y procesador de textos ───────────────────────────────────────────────

# El mensaje cuando falta playwright. Dice los DOS pasos porque son DOS: el
# pip install trae la libreria pero NO el navegador, y el fallo del segundo
# paso (Executable doesn't exist) es el que mas tiempo hace perder. Y dice
# donde NO esta instalado: playwright no vive en el venv del producto
# (~/.cognia/venv), asi que este camino corre en el repo y falla en una
# instalacion limpia -- fingir lo contrario seria mentir en el unico sitio
# donde el duenio va a mirar.
_FALTA_PLAYWRIGHT = (
    "el PDF automatico necesita playwright, que NO viene con Cognia. Son DOS "
    "pasos, no uno:\n"
    "  1) pip install playwright\n"
    "  2) playwright install chromium\n"
    "(el paso 1 instala la libreria; el 2 baja el navegador. Sin el 2 el error "
    "es \"Executable doesn't exist\".)\n"
    "OJO: playwright no esta en el venv del producto (~/.cognia/venv), asi que "
    "esto funciona desde el repo y falla en una instalacion limpia.\n"
    "EL CAMINO QUE SIEMPRE FUNCIONA, sin instalar nada: abre el cuaderno "
    "(/grabar-clase ver), boton Imprimir -> Guardar como PDF. La pagina ya "
    "viene preparada para el papel.")


@contextmanager
def _pagina(html_ruta: Path, ancho: int = 1240):
    """Un HTML ya escrito, abierto en un Chromium headless.

    El import es PEREZOSO: playwright es un extra, y un import arriba haria
    que TODO el modulo -- incluido el cuaderno HTML, que no lo necesita --
    dejara de importarse donde no esta.

    Es un context manager y no una funcion que devuelve tres cosas porque son
    DOS recursos que hay que cerrar en orden inverso (el navegador y el
    proceso de playwright) y cualquier fallo por el medio -- que aqui es lo
    normal: falta chromium, el HTML no abre -- dejaria un chromium colgado.
    """
    try:
        from playwright.sync_api import sync_playwright
    except Exception as exc:
        raise ErrorExportacion("%s\n(detalle: %s: %s)"
                               % (_FALTA_PLAYWRIGHT, type(exc).__name__, exc))
    pw = sync_playwright().start()
    nav = None
    try:
        try:
            nav = pw.chromium.launch()
        except Exception as exc:
            raise ErrorExportacion("%s\n(detalle al arrancar chromium: %s: %s)"
                                   % (_FALTA_PLAYWRIGHT, type(exc).__name__, exc))
        pag = nav.new_page(viewport={"width": ancho, "height": 1600})
        pag.goto(html_ruta.as_uri(), wait_until="load")
        yield pag
    finally:
        if nav is not None:
            try:
                nav.close()
            except Exception as exc:
                log.warning("clases.vista: no pude cerrar chromium: %s", exc)
        try:
            pw.stop()
        except Exception as exc:
            log.warning("clases.vista: no pude parar playwright: %s", exc)


def _preparar_papel(pag) -> dict:
    """Deja la pagina como la deja el boton Imprimir, y espera a las imagenes.

    Se llama a la MISMA funcion de la pagina (`__prepararImpresion`) que el
    evento beforeprint: `page.pdf()` no dispara beforeprint, y reimplementar
    aqui lo que hace el boton daria dos PDF distintos segun quien lo pida.

    Lo del loading="eager" es un SEGURO, y hay que decir hasta donde llega
    porque se midio: en un cuaderno de 26 fotos de pizarra con el scroller a
    20.272 px de alto, Chromium carga las 26 ANTES de tocar nada, aunque
    todas lleven loading="lazy". El motivo es que aqui las imagenes son data:
    URI y no hay nada que aplazar -- el aplazamiento existe para ahorrar red.
    O sea que con el cuaderno tal y como esta HOY, quitar el eager no cambia
    el PDF: no hay test que falle sin el, y esta dicho a proposito en vez de
    vender un arreglo que no arregla nada medible.

    Se deja igual porque el dia que un adjunto se sirva por file:// o por
    http (una foto que no cabe embebida, otro navegador con otra politica) la
    imagen que nunca entro en pantalla sale EN BLANCO en el papel, y en un PDF
    eso no se nota hasta que lo abre el que lo recibio. Cuesta una linea y la
    espera de abajo lo cierra: si el eager dispara una carga, aqui se espera a
    que termine.
    """
    info = pag.evaluate("() => window.__prepararImpresion()") or {}
    pag.evaluate("""() => Promise.all(
        Array.from(document.querySelectorAll('img.adj')).map(
            i => i.complete ? null : new Promise(r => {
                i.addEventListener('load', r); i.addEventListener('error', r);
            })))""")
    return dict(info)


@contextmanager
def _html_de_origen(materias=None, origen=None, ahora=None):
    """El HTML que se va a convertir, ya en disco.

    Si dan `origen` se usa ese fichero tal cual (y NO se borra: es del
    duenio). Si no, se genera uno en un temporal que se borra al salir --
    dejarlo suelto en ~/.cognia crearia un cuaderno.html que nadie pidio y que
    ademas pisaria al de `export()`.
    """
    if origen:
        ruta = Path(origen).expanduser()
        if not ruta.is_file():
            raise ErrorExportacion("no existe el HTML de origen: %s" % ruta)
        yield ruta, ruta
        return
    import shutil
    import tempfile
    tmp = tempfile.mkdtemp(prefix="cognia_export_")
    try:
        ruta = Path(tmp) / "cuaderno.html"
        sufijo = (" · " + ", ".join(materias)) if materias else ""
        ruta.write_text(render_html(construir(materias, ahora=ahora),
                                    "Cuaderno de clase" + sufijo),
                        encoding="utf-8")
        yield ruta, None
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def export_pdf(path=None, materias=None, origen=None, ahora=None) -> dict:
    """El cuaderno en PDF con playwright. EXTRA OPCIONAL (ver _FALTA_PLAYWRIGHT).

    Devuelve {'pdf', 'html', 'clips', 'imagenes', 'enlazados'}. 'html' es el
    fichero de origen SOLO si lo dio el llamante (el generado aqui es
    temporal y ya no existe al volver). Los clips de audio y los adjuntos que
    solo van como enlace se CUENTAN y se declaran en el pie del papel, porque
    son exactamente lo que un PDF no puede llevar.

    `origen` permite pasar un HTML ya escrito (p.ej. uno de los cuadernos por
    asignatura) en vez de generar otro.
    """
    destino = Path(path).expanduser() if path else (
        Path.home() / ".cognia" / "cuaderno.pdf")
    destino.parent.mkdir(parents=True, exist_ok=True)
    with _html_de_origen(materias, origen, ahora) as (ruta_html, suyo):
        with _pagina(ruta_html) as pag:
            info = _preparar_papel(pag)
            pag.pdf(path=str(destino), format="A4", print_background=True,
                    margin={"top": "12mm", "bottom": "12mm",
                            "left": "12mm", "right": "12mm"})
    return {"pdf": destino, "html": suyo,
            "clips": int(info.get("clips") or 0),
            "imagenes": int(info.get("imagenes") or 0),
            "enlazados": int(info.get("enlazados") or 0)}


def export_dom(path=None, materias=None, origen=None, ahora=None) -> Path:
    """El DOM YA RENDERIZADO, para subir a un procesador de textos.

    POR QUE NO VALE EL HTML CRUDO. MEDIDO: subir el cuaderno tal cual a Google
    Docs da un documento VACIO de 272 caracteres. No es un bug de Google: esta
    pagina pinta TODO su contenido con JavaScript desde el JSON embebido (ver
    la cabecera: es lo que la hace segura), y un importador de documentos no
    ejecuta JS -- solo ve el esqueleto. Lo que si se convierte es el DOM
    despues de que el JS corra, que es lo que escribe esta funcion.

    REPRODUCIDO aqui contando el texto que queda al quitar <script>, <style> y
    <template> -- que es lo que ve un importador que no ejecuta JS -- sobre un
    cuaderno de dos clases de Fisica: el HTML crudo son 32.895 caracteres de
    fichero pero solo 265 de texto importable (del orden de los 272 del
    documento vacio que salio en Google Docs), y el DOM ya renderizado, 1.546.
    El fichero nunca estuvo vacio: lo estaba lo que el importador sabe leer.

    Necesita playwright por el mismo motivo que el PDF: hace falta un
    navegador de verdad para que el JS corra.
    """
    destino = Path(path).expanduser() if path else (
        Path.home() / ".cognia" / "cuaderno_para_documento.html")
    destino.parent.mkdir(parents=True, exist_ok=True)
    with _html_de_origen(materias, origen, ahora) as (ruta_html, _suyo):
        with _pagina(ruta_html) as pag:
            # Las transcripciones plegadas: un <details> cerrado SI esta en el
            # DOM (a diferencia de la impresion), pero el importador lo aplana
            # raro. Se abren para que el documento lleve el texto suelto.
            _preparar_papel(pag)
            destino.write_text(pag.content(), encoding="utf-8")
    return destino


_FALTA_DOCX = ("el .docx directo necesita python-docx, que no viene con "
               "Cognia:\n  pip install python-docx\n"
               "(no esta en el venv del producto, ~/.cognia/venv). Sin el, el "
               "camino para llevar el cuaderno a Word o a Google Docs es "
               "export_dom(): sube ESE fichero, no el HTML crudo.")

# Como se pinta cada tipo de bloque del documento en el .docx. Es el PUNTO DE
# EXTENSION: aniadir un tipo es aniadir una fila, no tocar el bucle. El valor
# es (estilo de parrafo, prefijo).
#
# Las claves son los `documento.TIPO_*` escritos como literal y no importados,
# porque `clases.documento` se importa PEREZOSAMENTE (arriba obligaria a tener
# el paquete entero solo para pintar un HTML). Un tipo nuevo que no este aqui
# no se pierde: cae en ("Normal", "") y se escribe igual.
_ESTILO_BLOQUE = {
    "titulo": ("Heading 1", ""),
    "subtitulo": ("Heading 2", ""),
    "parrafo": ("Normal", ""),
    "lista": ("List Bullet", ""),
    "cita": ("Intense Quote", ""),
    "deber": ("List Bullet", "Deberes: "),
    "duda": ("List Bullet", "Duda: "),
    "examen": ("List Bullet", "Entra en el examen: "),
    "tabla": ("Normal", ""),
}


def _ruta_de_adjunto(nombre: str, jornada: str = "") -> Path:
    """Donde esta un adjunto del documento. Se busca en la jornada que dice su
    meta y, si no, por todas: un PNG de formula lo escribio la jornada en la
    que se genero, y el documento de la materia lo referencia por nombre."""
    for j in ([jornada] if jornada else []) + list(alm.jornadas()):
        if not j:
            continue
        try:
            ruta = alm.ruta_adjunto(j, nombre)
        except OSError as exc:
            log.warning("clases.vista: %s/%s ilegible: %s", j, nombre, exc)
            continue
        if ruta.is_file():
            return ruta
    return Path("")


def export_docx(materia: str, path=None):
    """El documento de UNA materia como .docx, desde sus BLOQUES.

    No pasa por el HTML: los bloques (`clases/documento.py`) ya son la
    estructura del documento -- titulo, parrafo, lista, formula, grafica --,
    asi que convertirlos a Word es directo y sale un fichero editable de
    verdad, con estilos, en vez de un volcado de pagina web. Las formulas y
    las graficas van como IMAGEN (es lo que son: un PNG que genero
    `clases/mates.py`) y lo que falte se escribe DICIENDO que falta, nunca en
    blanco.
    """
    try:
        import docx as _docx
        # Explicito: `import docx` no garantiza que docx.shared este cargado,
        # y descubrirlo en el add_picture seria un AttributeError a mitad de
        # documento en vez de un error con instrucciones aqui.
        from docx.shared import Cm as _Cm
    except Exception as exc:
        raise ErrorExportacion("%s\n(detalle: %s: %s)"
                               % (_FALTA_DOCX, type(exc).__name__, exc))
    try:
        from cognia.clases import documento as doc
    except Exception as exc:
        raise ErrorExportacion("no pude importar cognia.clases.documento "
                               "(%s: %s)" % (type(exc).__name__, exc))
    d = doc.abrir(materia, crear=False)
    destino = Path(path).expanduser() if path else (
        Path.home() / ".cognia" / ("%s.docx" % alm._seguro(materia)))
    destino.parent.mkdir(parents=True, exist_ok=True)

    docu = _docx.Document()
    docu.add_heading(materia, 0)
    if not d.bloques:
        docu.add_paragraph("Este documento todavia no tiene bloques: se llena "
                           "al generar los apuntes de una clase de esta "
                           "materia (/grabar-clase parar).")
    for b in d.bloques:
        if b.tipo in ("formula", "grafica", "imagen"):
            fichero = str(b.meta.get("adjunto") or b.meta.get("png") or "")
            ruta = _ruta_de_adjunto(fichero, str(b.meta.get("jornada") or ""))
            if fichero and ruta.is_file():
                try:
                    docu.add_picture(str(ruta), width=_Cm(14))
                except Exception as exc:
                    log.warning("clases.vista: no pude meter %s: %s", ruta, exc)
                    docu.add_paragraph("[no pude insertar la imagen '%s': %s]"
                                       % (fichero, exc))
            else:
                docu.add_paragraph("[falta la imagen '%s' de este %s]"
                                   % (fichero or "sin nombre", b.tipo))
            pie = (b.texto or "").strip() or str(b.meta.get("latex") or
                                                 b.meta.get("expresion") or "")
            if pie:
                # El pie va en un RUN y no en el parrafo: `Paragraph` no tiene
                # .italic (aceptaria el atributo y no haria nada, que es la
                # peor de las tres opciones posibles).
                docu.add_paragraph().add_run(pie).italic = True
            continue
        estilo, prefijo = _ESTILO_BLOQUE.get(b.tipo, ("Normal", ""))
        texto = prefijo + (b.texto or "").strip()
        if not texto.strip():
            continue
        try:
            docu.add_paragraph(texto, style=estilo)
        except KeyError:
            # Una plantilla de Word sin ese estilo no puede tragarse el
            # bloque: se escribe sin estilo antes que perderlo.
            log.warning("clases.vista: el .docx no tiene el estilo %r", estilo)
            docu.add_paragraph(texto)
    docu.save(str(destino))
    return destino


if __name__ == "__main__":
    # La puerta de linea de comandos del modulo. Cada camino de exportacion
    # tiene la suya para poder probarlo sin el REPL:
    #   python -m cognia.clases.vista                 un solo HTML
    #   python -m cognia.clases.vista --por-materia   un HTML por asignatura
    #   python -m cognia.clases.vista --pdf           PDF (necesita playwright)
    #   python -m cognia.clases.vista --dom           DOM post-JS para Docs/Word
    #   python -m cognia.clases.vista --docx <mat>    .docx (necesita python-docx)
    import sys
    abrir = "--no-open" not in sys.argv
    try:
        if "--por-materia" in sys.argv:
            res = export_materias(open_browser=abrir)
            print("indice -> %s" % res["indice"])
            for m, r in res["ficheros"].items():
                print("  %-24s %s" % (m, r.name))
            for a in res["avisos"]:
                print("  aviso: %s" % a)
        elif "--pdf" in sys.argv:
            res = export_pdf()
            print("pdf -> %s  (%d imagenes, %d clips de audio fuera del papel)"
                  % (res["pdf"], res["imagenes"], res["clips"]))
        elif "--dom" in sys.argv:
            print("dom -> %s" % export_dom())
        elif "--docx" in sys.argv:
            i = sys.argv.index("--docx")
            materia = sys.argv[i + 1] if len(sys.argv) > i + 1 else ""
            if not materia:
                raise SystemExit("uso: --docx <materia>")
            print("docx -> %s" % export_docx(materia))
        else:
            print("cuaderno -> %s" % export(open_browser=abrir))
    except ErrorExportacion as exc:
        # Un extra que falta no es un traceback: es un mensaje con los pasos.
        raise SystemExit(str(exc))
