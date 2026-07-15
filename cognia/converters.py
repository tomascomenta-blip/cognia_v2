# -*- coding: utf-8 -*-
"""Conversor universal de documentos a texto/markdown (MarkItDown nativo).

Mandato 2026-07-14: "MarkItDown -> conversor universal interno de
documentos." Cognia ya ingería texto plano/código y PDF (pdfplumber) en
ingest.py; este módulo centraliza la conversión y AMPLÍA los formatos, para
que cualquier subsistema (ingest, /resumir, cuaderno, agente) obtenga texto
limpio de un documento sin repetir la lógica.

Formatos:
- texto/código: read directo (utf-8, errores reemplazados).
- .pdf: pdfplumber (dep opcional, ya usada).
- .html/.htm: extracción con html.parser (stdlib) — quita script/style,
  colapsa espacios, conserva saltos de bloque. Mejor que el strip crudo por
  regex que hacía http_get (área navegador).
- .csv/.tsv: stdlib csv → filas separadas por ' | ' (tabla legible por LLM).
- .docx: python-docx si está (dep opcional); si no, mensaje accionable.
- .xlsx: openpyxl si está (dep opcional).
- .json: pretty-print estable (para que el chunker lo trate como texto).

Patrón de deps opcionales igual que pdfplumber: import perezoso dentro de
la rama; si falta, RuntimeError con el pip install exacto (no rompe el
import del módulo). Nada que baje modelos ni toque la red.
"""
import csv
import io
import json
from html.parser import HTMLParser
from pathlib import Path

# Extensiones que este conversor sabe transformar más allá de leer crudo.
CONVERTIBLES = {".pdf", ".html", ".htm", ".csv", ".tsv", ".docx", ".xlsx",
                ".json"}


class _TextoDeHTML(HTMLParser):
    """Extrae texto visible; ignora script/style; salto tras bloques."""
    _BLOQUE = {"p", "div", "br", "li", "tr", "h1", "h2", "h3", "h4", "h5",
               "h6", "section", "article", "header", "footer", "ul", "ol",
               "table", "blockquote", "pre"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self._partes = []
        self._saltar = 0

    def handle_starttag(self, tag, attrs):
        if tag in ("script", "style"):
            self._saltar += 1
        elif tag in self._BLOQUE:
            self._partes.append("\n")

    def handle_endtag(self, tag):
        if tag in ("script", "style") and self._saltar:
            self._saltar -= 1
        elif tag in self._BLOQUE:
            self._partes.append("\n")

    def handle_data(self, data):
        if not self._saltar and data.strip():
            self._partes.append(data.strip())

    def texto(self) -> str:
        crudo = " ".join(self._partes)
        # colapsar espacios pero preservar dobles saltos de bloque
        import re
        crudo = re.sub(r"[ \t]+", " ", crudo)
        crudo = re.sub(r"\s*\n\s*", "\n", crudo)
        crudo = re.sub(r"\n{3,}", "\n\n", crudo)
        return crudo.strip()


def html_a_texto(html: str) -> str:
    p = _TextoDeHTML()
    p.feed(html)
    return p.texto()


def _csv_a_texto(path: Path, delim: str) -> str:
    filas = []
    with path.open("r", encoding="utf-8", errors="replace", newline="") as f:
        for row in csv.reader(f, delimiter=delim):
            filas.append(" | ".join(cell.strip() for cell in row))
    return "\n".join(filas)


def _docx_a_texto(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        raise RuntimeError("python-docx no instalado: pip install python-docx")
    d = docx.Document(str(path))
    partes = [p.text for p in d.paragraphs if p.text.strip()]
    for tabla in d.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells]
            if any(celdas):
                partes.append(" | ".join(celdas))
    return "\n\n".join(partes)


def _xlsx_a_texto(path: Path) -> str:
    try:
        import openpyxl
    except ImportError:
        raise RuntimeError("openpyxl no instalado: pip install openpyxl")
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    partes = []
    for hoja in wb.worksheets:
        partes.append(f"# {hoja.title}")
        for fila in hoja.iter_rows(values_only=True):
            celdas = ["" if v is None else str(v) for v in fila]
            if any(c.strip() for c in celdas):
                partes.append(" | ".join(celdas))
    wb.close()
    return "\n".join(partes)


def convertir_a_texto(path) -> str:
    """Devuelve el texto de `path` según su extensión. Lanza RuntimeError con
    el pip install exacto si falta una dependencia opcional."""
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        try:
            import pdfplumber
        except ImportError:
            raise RuntimeError("pdfplumber no instalado: pip install pdfplumber")
        pages = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
        return "\n\n".join(pages)
    if ext in (".html", ".htm"):
        return html_a_texto(path.read_text(encoding="utf-8", errors="replace"))
    if ext == ".csv":
        return _csv_a_texto(path, ",")
    if ext == ".tsv":
        return _csv_a_texto(path, "\t")
    if ext == ".docx":
        return _docx_a_texto(path)
    if ext == ".xlsx":
        return _xlsx_a_texto(path)
    if ext == ".json":
        try:
            obj = json.loads(path.read_text(encoding="utf-8", errors="replace"))
            return json.dumps(obj, indent=2, ensure_ascii=False)
        except Exception:
            return path.read_text(encoding="utf-8", errors="replace")
    return path.read_text(encoding="utf-8", errors="replace")
